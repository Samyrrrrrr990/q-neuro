"""Build synchronized modular LaTeX and a publication-style Word manuscript."""

from __future__ import annotations

import hashlib
import json
import re
import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper"
SOURCE = PAPER / "source"
TABLES = PAPER / "tables"
FIGURE_SOURCE = ROOT / "research" / "figures" / "generated"
FIGURES = PAPER / "figures"

TITLE = "Exact Real Controls Overturn an Apparent Complex-Valued Robustness Advantage"
SUBTITLE = "A preregistered falsification study across synthetic sequential task families"
AUTHOR = "Samyar Shafiee"
AFFILIATION = "Independent Researcher · Toronto, Canada"
VERSION = "Research manuscript · v1.0.0 · 14 August 2026"
RUNNING_TITLE = "Q-Neuro · Exact controls and comparator-sensitive conclusions"

NAVY = "193552"
BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
TEAL = "188F87"
CORAL = "D95C4B"
MIST = "F4F6F9"
PALE_TEAL = "E5F3F1"
PALE_CORAL = "F8ECE9"
GRAY = "5E6C78"

MODULES = [
    "abstract",
    "introduction",
    "related_work",
    "theory",
    "architecture",
    "training",
    "neuroworld",
    "experiments",
    "results",
    "ablations",
    "interpretability",
    "limitations",
    "safety",
    "discussion",
    "conclusion",
    "appendices",
]

CITATION_RE = re.compile(r"\[@([^\]]+)\]")
INLINE_RE = re.compile(r"(\*\*.*?\*\*|`.*?`)")


@dataclass
class Block:
    kind: str
    text: str = ""
    meta: tuple[str, ...] = ()


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def normalize_docx_container(path: Path) -> None:
    """Rewrite the OPC ZIP with stable ordering and timestamps."""
    temporary = path.with_suffix(".normalized.docx")
    fixed_time = (2026, 8, 14, 0, 0, 0)
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary, "w") as target:
        for member in sorted(source.infolist(), key=lambda item: item.filename):
            info = zipfile.ZipInfo(member.filename, date_time=fixed_time)
            info.compress_type = member.compress_type
            info.create_system = member.create_system
            info.external_attr = member.external_attr
            info.internal_attr = member.internal_attr
            target.writestr(info, source.read(member.filename), compresslevel=9)
    temporary.replace(path)


def parse_source(path: Path) -> list[Block]:
    blocks: list[Block] = []
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            blocks.append(Block("paragraph", " ".join(value.strip() for value in paragraph)))
            paragraph.clear()

    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            flush()
            continue
        if line.startswith("{{") and line.endswith("}}"):
            flush()
            payload = line[2:-2]
            kind, _, value = payload.partition(":")
            blocks.append(Block(kind, meta=tuple(value.split("|"))))
            continue
        if line.startswith("### "):
            flush()
            blocks.append(Block("heading3", line[4:]))
            continue
        if line.startswith("## "):
            flush()
            blocks.append(Block("heading2", line[3:]))
            continue
        if line.startswith("# "):
            flush()
            blocks.append(Block("heading1", line[2:]))
            continue
        if line.startswith("- "):
            flush()
            blocks.append(Block("bullet", line[2:]))
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            flush()
            blocks.append(Block("numbered", numbered.group(1)))
            continue
        paragraph.append(line)
    flush()
    return blocks


def load_references() -> list[dict[str, str]]:
    return json.loads((PAPER / "references.json").read_text(encoding="utf-8"))


def citation_numbers(text: str, number_by_key: dict[str, int]) -> str:
    def replace(match: re.Match[str]) -> str:
        keys = [value.strip().lstrip("@") for value in match.group(1).split(";")]
        numbers = [number_by_key[key] for key in keys]
        return "[" + ",".join(str(number) for number in numbers) + "]"

    return CITATION_RE.sub(replace, text)


def plain_markdown(text: str, number_by_key: dict[str, int]) -> str:
    text = citation_numbers(text, number_by_key)
    return text.replace("**", "").replace("`", "")


def add_inline_runs(paragraph: Any, text: str, number_by_key: dict[str, int]) -> None:
    text = citation_numbers(text, number_by_key)
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(
    cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table: Any, widths_dxa: list[int], *, indent_dxa: int = 120) -> None:
    """Apply fixed Word table geometry with matching table, grid, and cell widths."""

    if len(widths_dxa) != len(table.columns):
        raise ValueError("table width count must match the number of columns")
    table.autofit = False
    table_pr = table._tbl.tblPr
    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths_dxa)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:w"), str(indent_dxa))
    table_indent.set(qn("w:type"), "dxa")
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa, strict=True):
            cell.width = Inches(width / 1440)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def configure_list_numbering(document: Document, style_name: str) -> None:
    """Apply the narrative-proposal list geometry to a built-in real numbering style."""

    style = document.styles[style_name]
    style.paragraph_format.left_indent = Inches(0.375)
    style.paragraph_format.first_line_indent = Inches(-0.194)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.208
    number_properties = style.element.pPr.numPr
    number_id = str(number_properties.numId.val)
    numbering = document.part.numbering_part.element
    number = numbering.find(f'./w:num[@w:numId="{number_id}"]', numbering.nsmap)
    abstract_id = str(number.find("./w:abstractNumId", numbering.nsmap).get(qn("w:val")))
    abstract = numbering.find(f'./w:abstractNum[@w:abstractNumId="{abstract_id}"]', numbering.nsmap)
    level = abstract.find('./w:lvl[@w:ilvl="0"]', numbering.nsmap)
    suffix = level.find("./w:suff", numbering.nsmap)
    if suffix is None:
        suffix = OxmlElement("w:suff")
        level.insert(1, suffix)
    suffix.set(qn("w:val"), "tab")
    paragraph_properties = level.find("./w:pPr", numbering.nsmap)
    tabs = paragraph_properties.find("./w:tabs", numbering.nsmap)
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        paragraph_properties.insert(0, tabs)
    tab = tabs.find("./w:tab", numbering.nsmap)
    if tab is None:
        tab = OxmlElement("w:tab")
        tabs.append(tab)
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    indent = paragraph_properties.find("./w:ind", numbering.nsmap)
    if indent is None:
        indent = OxmlElement("w:ind")
        paragraph_properties.append(indent)
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "279")


def prevent_row_split(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_repeat_table_header(row: Any) -> None:
    repeat_header(row)


def set_picture_alt(document: Document, alt_text: str) -> None:
    shape = document.inline_shapes[-1]
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", "Q-Neuro research figure")


def add_field(paragraph: Any, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_paragraph_bottom_border(paragraph: Any, color: str, size: int = 12) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DEEP_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    configure_list_numbering(document, "List Bullet")
    configure_list_numbering(document, "List Number")

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)

    subtitle = styles.add_style("Paper Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False

    equation = styles.add_style("Paper Equation", WD_STYLE_TYPE.PARAGRAPH)
    equation.font.name = "Cambria Math"
    equation.font.size = Pt(10.5)
    equation.font.color.rgb = RGBColor.from_string(NAVY)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.left_indent = Inches(0.25)
    equation.paragraph_format.right_indent = Inches(0.25)
    equation.paragraph_format.space_before = Pt(5)
    equation.paragraph_format.space_after = Pt(9)

    note = styles.add_style("Table Note", WD_STYLE_TYPE.PARAGRAPH)
    note.font.name = "Calibri"
    note.font.size = Pt(8.5)
    note.font.color.rgb = RGBColor.from_string(GRAY)
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(10)

    quote = styles["Intense Quote"]
    quote.font.name = "Calibri"
    quote.font.size = Pt(10.5)
    quote.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
    quote.paragraph_format.left_indent = Inches(0.3)
    quote.paragraph_format.right_indent = Inches(0.3)
    quote.paragraph_format.space_before = Pt(8)
    quote.paragraph_format.space_after = Pt(10)

    header = section.header.paragraphs[0]
    header.text = RUNNING_TITLE
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)
    set_paragraph_bottom_border(header, "D9E3EA", 4)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Synthetic, nonclinical research manuscript  ·  ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    add_field(footer, "PAGE", "1")

    settings = document.settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def add_cover(document: Document) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(42)

    badge = document.add_table(rows=1, cols=1)
    badge.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(badge.rows[0])
    set_table_geometry(badge, [4824], indent_dxa=0)
    cell = badge.cell(0, 0)
    set_cell_shading(cell, PALE_TEAL)
    set_cell_margins(cell, 70, 120, 70, 120)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("RESULTS ARTICLE  ·  FALSIFICATION FIRST")
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    title = document.add_paragraph(TITLE, style="Title")
    title.paragraph_format.space_before = Pt(30)

    subtitle = document.add_paragraph(SUBTITLE, style="Paper Subtitle")
    subtitle.paragraph_format.left_indent = Inches(0.45)
    subtitle.paragraph_format.right_indent = Inches(0.45)

    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(8)
    author.paragraph_format.space_after = Pt(3)
    run = author.add_run(AUTHOR)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    affiliation = document.add_paragraph(AFFILIATION)
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.paragraph_format.space_after = Pt(18)
    for run in affiliation.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(GRAY)

    version = document.add_paragraph(VERSION)
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.paragraph_format.space_after = Pt(28)
    for run in version.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(GRAY)

    boundary = document.add_table(rows=1, cols=1)
    boundary.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(boundary.rows[0])
    set_table_geometry(boundary, [8352], indent_dxa=0)
    cell = boundary.cell(0, 0)
    set_cell_shading(cell, PALE_CORAL)
    set_cell_margins(cell, 150, 180, 150, 180)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(
        "EVIDENCE BOUNDARY  ·  Synthetic experiments only. No patient data, clinical validity, "
        "physical quantum computation, or claim of universal model superiority."
    )
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(CORAL)

    metrics = document.add_paragraph()
    metrics.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metrics.paragraph_format.space_before = Pt(24)
    metrics.paragraph_format.space_after = Pt(0)
    run = metrics.add_run(
        "4,800 INDEPENDENT-TASK EFFECT CELLS   ·   1 FROZEN LAW TEST   ·   GRAND BENCHMARK SEALED"
    )
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(DEEP_BLUE)

    document.add_page_break()
    heading = document.add_paragraph("Contents", style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)
    intro = document.add_paragraph(
        "A structured map of the manuscript. Word's Navigation pane uses the same semantic "
        "heading hierarchy throughout the document."
    )
    intro.paragraph_format.space_after = Pt(14)
    contents = [
        ("01", "Abstract", "09", "Results"),
        ("02", "Introduction", "10", "Ablations"),
        ("03", "Related work", "11", "Interpretability and state audit"),
        ("04", "Mathematical framework", "12", "Limitations"),
        ("05", "Architecture and controls", "13", "Safety and responsible use"),
        ("06", "Training and selection", "14", "Discussion"),
        ("07", "Synthetic task systems", "15", "Conclusion"),
        ("08", "Experimental design", "16", "Extended methods and results"),
    ]
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row_index, (left_no, left_title, right_no, right_title) in enumerate(contents):
        row = table.add_row()
        prevent_row_split(row)
        for cell, number, title_text in (
            (row.cells[0], left_no, left_title),
            (row.cells[1], right_no, right_title),
        ):
            cell.width = Inches(3.25)
            set_cell_margins(cell, 120, 150, 120, 150)
            if row_index % 2 == 0:
                set_cell_shading(cell, MIST)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            number_run = paragraph.add_run(f"{number}  ")
            number_run.bold = True
            number_run.font.color.rgb = RGBColor.from_string(TEAL)
            title_run = paragraph.add_run(title_text)
            title_run.bold = True
            title_run.font.color.rgb = RGBColor.from_string(NAVY)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [4680, 4680])
    note = document.add_paragraph(
        "References follow the supplementary results. Figures and tables are numbered in order "
        "of appearance and are generated from registered artifacts."
    )
    note.style = "Table Note"
    note.paragraph_format.space_before = Pt(12)
    document.add_page_break()


def add_figure(document: Document, meta: tuple[str, ...], figure_number: int) -> None:
    if len(meta) != 3:
        raise ValueError(f"Figure directive requires name, caption, alt text: {meta}")
    name, caption, alt_text = meta
    image_path = FIGURES / f"{name}.png"
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    with Image.open(image_path) as image_file:
        width_px, height_px = image_file.size
    width = 6.45
    height = width * height_px / width_px
    if height > 4.15:
        height = 4.15
        width = height * width_px / height_px
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width), height=Inches(height))
    set_picture_alt(document, alt_text)
    caption_paragraph = document.add_paragraph(style="Caption")
    prefix = caption_paragraph.add_run(f"Figure {figure_number} | ")
    prefix.bold = True
    caption_paragraph.add_run(caption)


def add_table(document: Document, identifier: str, table_number: int) -> None:
    payload = json.loads((TABLES / f"{identifier}.json").read_text(encoding="utf-8"))
    caption = document.add_paragraph(style="Caption")
    caption.paragraph_format.keep_with_next = True
    prefix = caption.add_run(f"Table {table_number} | ")
    prefix.bold = True
    caption.add_run(payload["caption"])

    columns = payload["columns"]
    rows = payload["rows"]
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    first_width = 1.8 if len(columns) >= 6 else 2.15
    other_width = (6.5 - first_width) / (len(columns) - 1)
    for index, label in enumerate(columns):
        cell = table.rows[0].cells[index]
        cell.width = Inches(first_width if index == 0 else other_width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(label)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])

    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.width = Inches(first_width if index == 0 else other_width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index % 2:
                set_cell_shading(cell, MIST)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            )
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(NAVY)
    first_width_dxa = round(first_width * 1440)
    other_width_dxa = round(other_width * 1440)
    set_table_geometry(
        table,
        [first_width_dxa, *([other_width_dxa] * (len(columns) - 1))],
    )
    note = document.add_paragraph(payload["note"], style="Table Note")
    note.paragraph_format.keep_with_next = False


def add_blocks_to_docx(
    document: Document,
    modules: dict[str, list[Block]],
    number_by_key: dict[str, int],
) -> tuple[int, int]:
    figure_number = 0
    table_number = 0
    for module in MODULES:
        for block in modules[module]:
            if block.kind == "heading1":
                document.add_paragraph(block.text, style="Heading 1")
            elif block.kind == "heading2":
                document.add_paragraph(block.text, style="Heading 2")
            elif block.kind == "heading3":
                document.add_paragraph(block.text, style="Heading 3")
            elif block.kind == "paragraph":
                paragraph = document.add_paragraph()
                add_inline_runs(paragraph, block.text, number_by_key)
            elif block.kind == "bullet":
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.left_indent = Inches(0.375)
                paragraph.paragraph_format.first_line_indent = Inches(-0.194)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.208
                add_inline_runs(paragraph, block.text, number_by_key)
            elif block.kind == "numbered":
                paragraph = document.add_paragraph(style="List Number")
                paragraph.paragraph_format.left_indent = Inches(0.375)
                paragraph.paragraph_format.first_line_indent = Inches(-0.194)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.208
                add_inline_runs(paragraph, block.text, number_by_key)
            elif block.kind == "equation":
                if len(block.meta) != 2:
                    raise ValueError(
                        f"Equation directive requires LaTeX and Word text: {block.meta}"
                    )
                paragraph = document.add_paragraph(style="Paper Equation")
                paragraph.add_run(block.meta[1])
                set_paragraph_bottom_border(paragraph, "D9E3EA", 4)
            elif block.kind == "figure":
                figure_number += 1
                add_figure(document, block.meta, figure_number)
            elif block.kind == "table":
                table_number += 1
                add_table(document, block.meta[0], table_number)
            else:
                raise ValueError(f"Unknown manuscript block: {block.kind}")
    return figure_number, table_number


def add_references(document: Document, references: list[dict[str, str]]) -> None:
    document.add_paragraph("References", style="Heading 1")
    for index, reference in enumerate(references, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.28)
        paragraph.paragraph_format.first_line_indent = Inches(-0.28)
        paragraph.paragraph_format.space_after = Pt(5)
        number = paragraph.add_run(f"{index}. ")
        number.bold = True
        paragraph.add_run(reference["text"])


def build_docx(
    modules: dict[str, list[Block]], references: list[dict[str, str]]
) -> tuple[int, int]:
    number_by_key = {reference["key"]: index for index, reference in enumerate(references, start=1)}
    document = Document()
    configure_document(document)
    document.core_properties.title = f"{TITLE}: {SUBTITLE}"
    document.core_properties.subject = (
        "Comparator-sensitive falsification of a complex-valued robustness claim"
    )
    document.core_properties.author = AUTHOR
    document.core_properties.keywords = (
        "Q-Neuro; exact real control; complex-valued networks; negative results; reproducibility"
    )
    document.core_properties.comments = (
        "Research manuscript generated from registered Q-Neuro artifacts; synthetic evidence only."
    )
    add_cover(document)
    figure_count, table_count = add_blocks_to_docx(document, modules, number_by_key)
    add_references(document, references)
    output = PAPER / "qneuro.docx"
    document.save(output)
    normalize_docx_container(output)
    return figure_count, table_count


def escape_tex(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
        "≤": r"$\leq$",
        "≥": r"$\geq$",
        "±": r"$\pm$",
        "×": r"$\times$",
        "→": r"$\rightarrow$",
        "−": "--",
        "–": "--",
        "—": "---",
    }
    return "".join(replacements.get(character, character) for character in text)


def tex_inline(text: str) -> str:
    citations: list[str] = []

    def hold_citation(match: re.Match[str]) -> str:
        keys = [value.strip().lstrip("@") for value in match.group(1).split(";")]
        citations.append(r"\citep{" + ",".join(keys) + "}")
        return f"QQCITE{len(citations) - 1}QQ"

    text = CITATION_RE.sub(hold_citation, text)
    inline_code: list[str] = []

    def hold_code(match: re.Match[str]) -> str:
        inline_code.append(r"\texttt{" + escape_tex(match.group(1)) + "}")
        return f"QQCODE{len(inline_code) - 1}QQ"

    text = re.sub(r"`([^`]+)`", hold_code, text)
    bold: list[str] = []

    def hold_bold(match: re.Match[str]) -> str:
        bold.append(r"\textbf{" + escape_tex(match.group(1)) + "}")
        return f"QQBOLD{len(bold) - 1}QQ"

    text = re.sub(r"\*\*(.+?)\*\*", hold_bold, text)
    text = escape_tex(text)
    for index, value in enumerate(citations):
        text = text.replace(f"QQCITE{index}QQ", value)
    for index, value in enumerate(inline_code):
        text = text.replace(f"QQCODE{index}QQ", value)
    for index, value in enumerate(bold):
        text = text.replace(f"QQBOLD{index}QQ", value)
    return text


def module_to_tex(module: str, blocks: list[Block]) -> str:
    lines = ["% Generated from paper/source. Do not edit by hand."]
    list_environment: str | None = None

    def close_list() -> None:
        nonlocal list_environment
        if list_environment is not None:
            lines.append(rf"\end{{{list_environment}}}")
            lines.append("")
            list_environment = None

    for block in blocks:
        if block.kind not in {"bullet", "numbered"}:
            close_list()
        if block.kind == "heading1":
            if module == "abstract":
                lines.append(r"\begin{abstract}")
            elif module == "appendices":
                lines.append(r"\appendix")
            else:
                lines.append(rf"\section{{{tex_inline(block.text)}}}")
        elif block.kind == "heading2":
            command = "section" if module == "appendices" else "subsection"
            lines.append(rf"\{command}{{{tex_inline(block.text)}}}")
        elif block.kind == "heading3":
            lines.append(rf"\subsubsection{{{tex_inline(block.text)}}}")
        elif block.kind == "paragraph":
            lines.append(tex_inline(block.text))
            lines.append("")
        elif block.kind in {"bullet", "numbered"}:
            required_environment = "itemize" if block.kind == "bullet" else "enumerate"
            if list_environment != required_environment:
                close_list()
                lines.append(rf"\begin{{{required_environment}}}")
                list_environment = required_environment
            lines.append(rf"\item {tex_inline(block.text)}")
        elif block.kind == "equation":
            lines.extend([r"\begin{equation}", block.meta[0], r"\end{equation}", ""])
        elif block.kind == "figure":
            name, caption, _ = block.meta
            lines.extend(
                [
                    r"\begin{figure}[t]",
                    r"\centering",
                    rf"\includegraphics[width=0.98\linewidth]{{figures/{name}.pdf}}",
                    rf"\caption{{{tex_inline(caption)}}}",
                    rf"\label{{fig:{name}}}",
                    r"\end{figure}",
                    "",
                ]
            )
        elif block.kind == "table":
            lines.extend([rf"\input{{tables/{block.meta[0]}.tex}}", ""])
        else:
            raise ValueError(f"Unknown LaTeX block: {block.kind}")
    close_list()
    if module == "abstract":
        lines.append(r"\end{abstract}")
    return "\n".join(lines).rstrip() + "\n"


def write_latex(modules: dict[str, list[Block]], references: list[dict[str, str]]) -> None:
    for module, blocks in modules.items():
        (PAPER / f"{module}.tex").write_text(module_to_tex(module, blocks), encoding="utf-8")
    (PAPER / "references.bib").write_text(
        "\n\n".join(reference["bibtex"] for reference in references) + "\n", encoding="utf-8"
    )
    includes = "\n".join(rf"\input{{{module}}}" for module in MODULES if module != "abstract")
    main = rf"""% Generated by paper/build_manuscript.py.
\documentclass[10pt]{{article}}
\usepackage[margin=0.8in]{{geometry}}
\usepackage[T1]{{fontenc}}
\usepackage[utf8]{{inputenc}}
\usepackage{{lmodern}}
\usepackage{{microtype}}
\usepackage{{graphicx}}
\usepackage{{booktabs}}
\usepackage{{array}}
\usepackage{{amsmath,amssymb}}
\usepackage{{caption}}
\usepackage{{enumitem}}
\usepackage[numbers,sort&compress]{{natbib}}
\usepackage[hidelinks]{{hyperref}}
\usepackage{{xcolor}}
\definecolor{{qnavy}}{{HTML}}{{{NAVY}}}
\definecolor{{qblue}}{{HTML}}{{{BLUE}}}
\setlength{{\parindent}}{{0pt}}
\setlength{{\parskip}}{{0.55em}}
\setlist{{nosep,leftmargin=1.4em}}
\captionsetup{{font=small,labelfont=bf}}
\title{{\textbf{{{TITLE}}}\\[0.5em]\large {SUBTITLE}}}
\author{{{AUTHOR}\\\small {AFFILIATION}}}
\date{{{VERSION}}}
\begin{{document}}
\maketitle
\input{{abstract}}
\tableofcontents
\clearpage
{includes}
\clearpage
\bibliographystyle{{unsrtnat}}
\bibliography{{references}}
\end{{document}}
"""
    (PAPER / "main.tex").write_text(main, encoding="utf-8")


def stage_figures(modules: dict[str, list[Block]]) -> list[str]:
    names: list[str] = []
    for blocks in modules.values():
        for block in blocks:
            if block.kind == "figure":
                name = block.meta[0]
                if name in names:
                    raise ValueError(f"Duplicate paper figure: {name}")
                names.append(name)
    for name in names:
        for extension in ("png", "pdf"):
            source = FIGURE_SOURCE / f"{name}.{extension}"
            if not source.exists():
                raise FileNotFoundError(source)
            shutil.copy2(source, FIGURES / source.name)
    manifest = {
        "manuscript_date": "2026-08-14",
        "count": len(names),
        "figures": [
            {
                "name": name,
                "png_sha256": sha256(FIGURES / f"{name}.png"),
                "pdf_sha256": sha256(FIGURES / f"{name}.pdf"),
            }
            for name in names
        ],
    }
    (FIGURES / "manifest.json").write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    return names


def write_manifest(figure_count: int, table_count: int, references: list[dict[str, str]]) -> None:
    source_files = [SOURCE / f"{module}.md" for module in MODULES]
    words = sum(len(path.read_text(encoding="utf-8").split()) for path in source_files)
    payload = {
        "title": TITLE,
        "subtitle": SUBTITLE,
        "author": AUTHOR,
        "version": VERSION,
        "scope": "Synthetic research only; no clinical validation.",
        "source_word_count": words,
        "figure_count": figure_count,
        "table_count": table_count,
        "reference_count": len(references),
        "source_sha256": {path.name: sha256(path) for path in source_files},
        "docx_sha256": sha256(PAPER / "qneuro.docx"),
    }
    (PAPER / "MANUSCRIPT_METADATA.json").write_text(
        json.dumps(payload, indent=2) + "\n", encoding="utf-8"
    )


def main() -> None:
    FIGURES.mkdir(parents=True, exist_ok=True)
    modules = {module: parse_source(SOURCE / f"{module}.md") for module in MODULES}
    references = load_references()
    staged = stage_figures(modules)
    write_latex(modules, references)
    figure_count, table_count = build_docx(modules, references)
    if figure_count != len(staged):
        raise RuntimeError("DOCX and LaTeX figure counts diverged")
    write_manifest(figure_count, table_count, references)
    print(
        f"Built modular LaTeX and qneuro.docx with {figure_count} figures, "
        f"{table_count} tables, and {len(references)} references"
    )


if __name__ == "__main__":
    main()
